"""
CLI Tool for Lynx Crypto Converter
Milestone 1: Parse balance files from command line
"""

import argparse
import sys
import os

# Add src directory to Python path if not already there
src_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)))
if src_dir not in sys.path:
    sys.path.insert(0, src_dir)

from parser import BalanceParser
from tabulate import tabulate
import json
import webbrowser
import subprocess
import requests
from dotenv import load_dotenv


def print_banner():
    """Print CLI banner"""
    banner = """
╔══════════════════════════════════════════════════════╗
║     LYNX CRYPTO CONVERTER - CLI TOOL                 ║
╚══════════════════════════════════════════════════════╝
    """
    print(banner)


def format_balance_table(balances):
    """Format balances as a table"""
    if not balances:
        return "No balances found"
    
    table_data = []
    for idx, balance in enumerate(balances, 1):
        table_data.append([
            idx,
            f"${balance['value']:,.2f}",
            balance.get('currency_symbol', 'N/A'),
            balance['context'][:50] + '...' if len(balance['context']) > 50 else balance['context']
        ])
    
    return tabulate(
        table_data,
        headers=['#', 'Value', 'Symbol', 'Context'],
        tablefmt='grid'
    )


def format_summary_table(summary):
    """Format summary statistics as a table"""
    summary_data = [
        ['Total Values Found', summary['total_values_found']],
        ['Total Sum', f"${summary['total_sum']:,.2f}"],
        ['Minimum Value', f"${summary['min_value']:,.2f}"],
        ['Maximum Value', f"${summary['max_value']:,.2f}"],
        ['Average Value', f"${summary['avg_value']:,.2f}"]
    ]
    
    return tabulate(summary_data, headers=['Metric', 'Value'], tablefmt='grid')


def parse_command(args):
    """Handle parse command"""
    try:
        print(f"\n📄 Parsing file: {args.file}")
        print("=" * 60)
        
        parser = BalanceParser(args.file)
        balances = parser.parse()
        summary = parser.get_summary()
        
        print(f"\n✅ Successfully parsed {len(balances)} balance value(s)\n")
        
        # Show summary
        print("📊 SUMMARY STATISTICS")
        print(format_summary_table(summary))
        
        # Show detailed balances if requested
        if args.detailed:
            print("\n📋 DETAILED BALANCES")
            print(format_balance_table(balances))
        
        # Export to JSON if requested
        if args.output:
            export_data = {
                'balances': balances,
                'summary': summary
            }
            with open(args.output, 'w') as f:
                json.dump(export_data, f, indent=2)
            print(f"\n💾 Results exported to: {args.output}")
        
        return 0
    
    except FileNotFoundError as e:
        print(f"\n❌ Error: File not found - {e}")
        return 1
    
    except ValueError as e:
        print(f"\n❌ Error: Invalid file - {e}")
        return 1
    
    except Exception as e:
        print(f"\n❌ Error: {e}")
        return 1


def validate_command(args):
    """Handle validate command"""
    try:
        print(f"\n🔍 Validating file: {args.file}")
        print("=" * 60)
        
        parser = BalanceParser(args.file)
        balances = parser.parse()
        
        if len(balances) > 0:
            print("\n✅ File is valid!")
            print(f"   Found {len(balances)} balance value(s)")
            print(f"   Total amount: ${parser.get_total():,.2f}")
        else:
            print("\n⚠️  File is valid but no numeric values were found")
        
        return 0
    
    except Exception as e:
        print(f"\n❌ Validation failed: {e}")
        return 1


def convert_command(args):
    """Handle convert command"""
    print(f"\n🔄 Converting file: {args.file}")
    print("=" * 60)
    
    api_url = "http://localhost:5001/api/convert"
    health_url = "http://localhost:5001/health"
    
    try:
        # Check if server is running
        response = requests.get(health_url, timeout=3)
        if response.status_code != 200:
            print("❌ API server is not running")
            print("💡 Start the server first with: python app.py")
            return 1
        
        print("✅ API server is running")
        
        # Check if file exists
        import os
        if not os.path.exists(args.file):
            print(f"❌ File not found: {args.file}")
            return 1
        
        # Send file to API
        with open(args.file, 'rb') as f:
            files = {'file': f}
            data = {}
            if args.currency:
                data['target_currency'] = args.currency
            
            print(f"🚀 Sending file to API for conversion...")
            response = requests.post(api_url, files=files, data=data, timeout=30)
        
        if response.status_code == 200:
            result = response.json()
            print("\n✅ Conversion completed successfully!")
            print(json.dumps(result, indent=2))
        else:
            error = response.json() if response.headers.get('content-type') == 'application/json' else {'error': response.text}
            print(f"\n❌ Conversion failed: {error.get('error', 'Unknown error')}")
            return 1
            
    except requests.exceptions.ConnectionError:
        print("❌ API server is not running")
        print("💡 Start the server first with: python app.py")
        return 1
    except Exception as e:
        print(f"❌ Error: {e}")
        return 1
    
    return 0


def api_command(args):
    """Handle api command"""
    print("\n🌐 Opening API Documentation...")
    print("=" * 60)
    
    api_url = "http://localhost:5001/"
    health_url = "http://localhost:5001/health"
    
    try:
        # Check if server is running
        response = requests.get(health_url, timeout=3)
        if response.status_code == 200:
            print("✅ API server is running")
            
            # Try to open in browser
            if webbrowser.open(api_url):
                print(f"🌐 Opened API documentation in browser: {api_url}")
            else:
                print(f"🌐 Please open in browser: {api_url}")
                
            print(f"📋 JSON API Docs: {api_url}api/docs")
            print(f"❤️  Health Check: {health_url}")
            
        else:
            print("❌ API server responded with error")
            return 1
            
    except requests.exceptions.ConnectionError:
        print("❌ API server is not running")
        print("💡 Start the server first with: python app.py")
        print("💡 Or use the launcher: ./lynx-launcher.sh")
        return 1
    except Exception as e:
        print(f"❌ Error connecting to API: {e}")
        return 1
    
    return 0


def setup_private_key():
    """Ensure private key is available"""
    home_dir = os.path.expanduser('~')
    key_dir = os.path.join(home_dir, 'Documents', 'key')
    key_file = os.path.join(key_dir, 'wallet.txt')
    
    if not os.path.exists(key_file):
        print("\n🔑 Private key not found. Please set up your wallet private key.")
        print(f"📁 Expected location: {key_file}")
        print("💡 Create the directory and file manually, or use the desktop installer.")
        return False
    return True


def send_command(args):
    """Handle send command - Direct blockchain transaction using real Web3 engine"""
    import json
    import asyncio
    from transaction_service import get_transaction_service
    from converter import crypto_converter
    
    # Load environment variables
    load_dotenv()
    
    # Check private key setup
    if not setup_private_key():
        return 1
    
    print(f"\n💸 Converting and sending to wallet: {args.file}")
    print("=" * 60)
    
    try:
        # Check if file exists
        if not os.path.exists(args.file):
            print(f"❌ File not found: {args.file}")
            return 1
        
        # First, convert the file to get/save conversion data
        print("🔄 Converting balances...")
        conversion_result = crypto_converter.convert_balances(args.file)
        
        if not conversion_result.get('success'):
            print(f"❌ Conversion failed: {conversion_result.get('error', 'Unknown error')}")
            return 1
        
        print(f"✅ Converted ${conversion_result['total_usd_amount']:,.2f} USD")
        
        # Show conversion summary
        print("\n💰 CONVERTED AMOUNTS:")
        for currency, amount in conversion_result['conversions'].items():
            print(f"   {currency}: {amount:.8f}")
        
        # Get transaction service
        transaction_service = get_transaction_service()
        
        # Check wallet setup
        if not transaction_service.account:
            print("\n❌ Wallet setup required!")
            print("💡 Run: ./setup-wallet.sh to configure your wallet")
            return 1
        
        if not transaction_service.web3 or not transaction_service.web3.is_connected():
            print("\n❌ Not connected to Ethereum network")
            print("💡 Check your ETH_NODE_URL in .env file")
            return 1
        
        # Get destination wallet address
        client_address = os.getenv('EURC_WALLET') or args.wallet_id
        if not client_address:
            print("\n❌ No destination wallet address configured")
            print("💡 Set EURC_WALLET in .env or use --wallet-id")
            return 1
        
        print(f"\n🎯 Destination: {client_address}")
        print("\n📤 SENDING TRANSACTIONS:")
        
        # Send each converted amount using real blockchain transactions
        wallet_transactions = []
        for currency, amount in conversion_result['conversions'].items():
            # Only send supported currencies (ETH, USDT, USDC)
            if currency.upper() not in ['ETH', 'USDT', 'USDC']:
                print(f"   ⏭️  {currency}: {amount:.8f} (Not supported for blockchain transactions)")
                continue
            
            print(f"   🚀 Sending {amount:.8f} {currency}...")
            
            try:
                # Use real blockchain transaction
                result = asyncio.run(transaction_service.send_eth(
                    to_address=client_address,
                    amount_eth=amount,
                    currency=currency
                ))
                
                if 'error' in result:
                    print(f"   ❌ {currency}: {result['error']}")
                    wallet_transactions.append({
                        'success': False,
                        'currency': currency,
                        'amount': amount,
                        'error': result['error'],
                        'wallet_address': client_address
                    })
                else:
                    print(f"   ✅ {amount:.8f} {currency} → {client_address[:10]}...")
                    if result.get('tx_hash'):
                        print(f"      🔗 TX: {result['tx_hash']}")
                    
                    wallet_transactions.append({
                        'success': True,
                        'currency': currency,
                        'amount': amount,
                        'tx_hash': result.get('tx_hash'),
                        'status': result.get('status', 'pending'),
                        'wallet_address': client_address
                    })
                    
            except Exception as e:
                print(f"   ❌ {currency}: Transaction failed - {str(e)}")
                wallet_transactions.append({
                    'success': False,
                    'currency': currency,
                    'amount': amount,
                    'error': str(e),
                    'wallet_address': client_address
                })
        
        # Summary
        successful_txs = [tx for tx in wallet_transactions if tx.get('success', False)]
        failed_txs = [tx for tx in wallet_transactions if not tx.get('success', True)]
        
        print(f"\n📊 SUMMARY:")
        print(f"   ✅ Successful: {len(successful_txs)}")
        print(f"   ❌ Failed: {len(failed_txs)}")
        
        if successful_txs:
            print("\n✅ Successfully sent converted amounts using real blockchain transactions!")
            return 0
        else:
            print("\n❌ All transactions failed")
            return 1
            
    except Exception as e:
        print(f"❌ Error: {e}")
        return 1


def demo_command(args):
    """Handle demo command"""
    print("\n🎬 DEMO MODE - Creating sample balance file...")
    print("=" * 60)
    
    try:
        from docx import Document
        
        # Create sample document
        doc = Document()
        doc.add_heading('Account Balances - November 2024', 0)
        
        doc.add_paragraph('Checking Account: $5,250.00')
        doc.add_paragraph('Savings Account: $12,800.50')
        doc.add_paragraph('Investment Portfolio: $45,000.00')
        doc.add_paragraph('Emergency Fund: $8,500.00')
        doc.add_paragraph('Crypto Wallet: $3,275.25')
        
        filename = 'demo_balances.docx'
        doc.save(filename)
        
        print(f"✅ Created sample file: {filename}")
        
        # Parse it
        print(f"\n📄 Parsing demo file...")
        parser = BalanceParser(filename)
        balances = parser.parse()
        summary = parser.get_summary()
        
        print("\n📊 DEMO RESULTS")
        print(format_summary_table(summary))
        
        print("\n📋 DEMO BALANCES")
        print(format_balance_table(balances))
        
        print(f"\n✅ Demo completed! You can now test with: python cli.py parse {filename}")
        
        return 0
    
    except Exception as e:
        print(f"\n❌ Demo failed: {e}")
        return 1


def list_conversions_command(args):
    """Handle list-conversions command"""
    print("\n📋 Saved Conversions")
    print("=" * 60)
    
    api_url = "http://localhost:5001/api/list-conversions"
    health_url = "http://localhost:5001/health"
    
    try:
        response = requests.get(health_url, timeout=3)
        if response.status_code != 200:
            print("❌ API server is not running")
            return 1
        
        response = requests.get(api_url, timeout=10)
        
        if response.status_code == 200:
            result = response.json()
            conversions = result.get('conversions', [])
            
            if not conversions:
                print("📄 No saved conversions found")
                return 0
            
            print(f"\n✅ Found {len(conversions)} saved conversion(s):\n")
            
            table_data = []
            for conv in conversions:
                status = "✅ Sent" if conv.get('sent') else "⏳ Pending"
                currencies = ", ".join(conv.get('currencies', []))
                timestamp = conv.get('timestamp', '')[:19].replace('T', ' ')
                
                table_data.append([
                    conv['id'],
                    f"${conv.get('total_usd', 0):,.2f}",
                    currencies,
                    timestamp,
                    status
                ])
            
            print(tabulate(
                table_data,
                headers=['Conversion ID', 'USD Amount', 'Currencies', 'Created', 'Status'],
                tablefmt='grid'
            ))
            
        else:
            print("❌ Failed to list conversions")
            return 1
            
    except requests.exceptions.ConnectionError:
        print("❌ API server is not running")
        return 1
    except Exception as e:
        print(f"❌ Error: {e}")
        return 1
    
    return 0


def send_saved_command(args):
    """Handle send-saved command - Direct blockchain transaction using real Web3 engine"""
    import json
    import asyncio
    from transaction_service import get_transaction_service
    from conversion_storage import conversion_storage
    
    # Load environment variables
    load_dotenv()
    
    # Check private key setup
    if not setup_private_key():
        return 1
    
    print(f"\n💸 Sending saved conversion: {args.conversion_id}")
    print("=" * 60)
    
    try:
        # Load saved conversion from conversions.json
        conversion = conversion_storage.get_conversion(args.conversion_id)
        
        if not conversion:
            print(f"❌ Conversion {args.conversion_id} not found")
            return 1
        
        if conversion.get('sent', False):
            print(f"❌ Conversion {args.conversion_id} already sent")
            return 1
        
        print(f"✅ Found saved conversion: ${conversion['total_usd_amount']:,.2f} USD")
        
        # Show conversion amounts
        print("\n💰 CONVERTED AMOUNTS:")
        for currency, amount in conversion['conversions'].items():
            print(f"   {currency}: {amount:.8f}")
        
        # Get transaction service
        transaction_service = get_transaction_service()
        
        # Check wallet setup
        if not transaction_service.account:
            print("\n❌ Wallet setup required!")
            print("💡 Run: ./setup-wallet.sh to configure your wallet")
            return 1
        
        if not transaction_service.web3 or not transaction_service.web3.is_connected():
            print("\n❌ Not connected to Ethereum network")
            print("💡 Check your ETH_NODE_URL in .env file")
            return 1
        
        # Get destination wallet address
        client_address = os.getenv('EURC_WALLET') or args.wallet_id
        if not client_address:
            print("\n❌ No destination wallet address configured")
            print("💡 Set EURC_WALLET in .env or use --wallet-id")
            return 1
        
        print(f"\n🎯 Destination: {client_address}")
        print("\n📤 SENDING TRANSACTIONS:")
        
        # Send each converted amount using real blockchain transactions
        wallet_transactions = []
        for currency, amount in conversion['conversions'].items():
            # Only send supported currencies (ETH, USDT, USDC)
            if currency.upper() not in ['ETH', 'USDT', 'USDC']:
                print(f"   ⏭️  {currency}: {amount:.8f} (Not supported for blockchain transactions)")
                continue
            
            print(f"   🚀 Sending {amount:.8f} {currency}...")
            
            try:
                # Use real blockchain transaction
                result = asyncio.run(transaction_service.send_eth(
                    to_address=client_address,
                    amount_eth=amount,
                    currency=currency
                ))
                
                if 'error' in result:
                    print(f"   ❌ {currency}: {result['error']}")
                    wallet_transactions.append({
                        'success': False,
                        'currency': currency,
                        'amount': amount,
                        'error': result['error'],
                        'wallet_address': client_address
                    })
                else:
                    print(f"   ✅ {amount:.8f} {currency} → {client_address[:10]}...")
                    if result.get('tx_hash'):
                        print(f"      🔗 TX: {result['tx_hash']}")
                    
                    wallet_transactions.append({
                        'success': True,
                        'currency': currency,
                        'amount': amount,
                        'tx_hash': result.get('tx_hash'),
                        'status': result.get('status', 'pending'),
                        'wallet_address': client_address
                    })
                    
            except Exception as e:
                print(f"   ❌ {currency}: Transaction failed - {str(e)}")
                wallet_transactions.append({
                    'success': False,
                    'currency': currency,
                    'amount': amount,
                    'error': str(e),
                    'wallet_address': client_address
                })
        
        # Mark as sent if any transactions succeeded
        successful_txs = [tx for tx in wallet_transactions if tx.get('success', False)]
        if successful_txs:
            conversion_storage.mark_as_sent(args.conversion_id)
        
        # Summary
        failed_txs = [tx for tx in wallet_transactions if not tx.get('success', True)]
        
        print(f"\n📊 SUMMARY:")
        print(f"   ✅ Successful: {len(successful_txs)}")
        print(f"   ❌ Failed: {len(failed_txs)}")
        
        if successful_txs:
            print("\n✅ Successfully sent saved conversion using real blockchain transactions!")
            return 0
        else:
            print("\n❌ All transactions failed")
            return 1
            
    except Exception as e:
        print(f"❌ Error: {e}")
        return 1


def main():
    """Main CLI entry point"""
    print_banner()
    
    parser = argparse.ArgumentParser(
        description='Lynx Crypto Converter - Balance File Parser',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  Parse a balance file:
    python cli.py parse balances.docx
  
  Parse with detailed output:
    python cli.py parse balances.docx --detailed
  
  Parse and export to JSON:
    python cli.py parse balances.docx --output results.json
  
  Validate a file:
    python cli.py validate balances.docx
  
  Run demo:
    python cli.py demo
  
  Convert balance file:
    python cli.py convert balances.docx
  
  Convert to specific currency:
    python cli.py convert balances.docx --currency EUR
  
  Send converted amounts to wallet:
    python cli.py send balances.docx
  
  Send to specific wallet ID:
    python cli.py send balances.docx --wallet-id custom_wallet_123
  
  List saved conversions:
    python cli.py list-conversions
  
  Send a saved conversion:
    python cli.py send-saved conv_20241118_143022_123456
  
  Open API documentation:
    python cli.py api
        """
    )
    
    subparsers = parser.add_subparsers(dest='command', help='Available commands')
    
    # Parse command
    parse_parser = subparsers.add_parser('parse', help='Parse balance file')
    parse_parser.add_argument('file', help='Path to balance file (.docx or .dox)')
    parse_parser.add_argument('-d', '--detailed', action='store_true', help='Show detailed balance list')
    parse_parser.add_argument('-o', '--output', help='Export results to JSON file')
    
    # Validate command
    validate_parser = subparsers.add_parser('validate', help='Validate balance file')
    validate_parser.add_argument('file', help='Path to balance file')
    
    # Demo command
    demo_parser = subparsers.add_parser('demo', help='Run demo with sample data')
    
    # Convert command
    convert_parser = subparsers.add_parser('convert', help='Convert balance file using API')
    convert_parser.add_argument('file', help='Path to balance file (.docx or .dox)')
    convert_parser.add_argument('-c', '--currency', help='Target currency (default: USD)')
    
    # Send command
    send_parser = subparsers.add_parser('send', help='Convert and send to wallet')
    send_parser.add_argument('file', help='Path to balance file (.docx or .dox)')
    send_parser.add_argument('-w', '--wallet-id', help='Wallet ID (defaults to client address)')
    
    # API command
    api_parser = subparsers.add_parser('api', help='Open API documentation')
    
    # List conversions command
    list_parser = subparsers.add_parser('list-conversions', help='List saved conversions')
    list_parser.add_argument('--pending-only', action='store_true', help='Show only pending conversions')
    
    # Send saved conversion command
    send_saved_parser = subparsers.add_parser('send-saved', help='Send a saved conversion to wallet')
    send_saved_parser.add_argument('conversion_id', help='ID of saved conversion to send')
    send_saved_parser.add_argument('-w', '--wallet-id', help='Wallet ID (defaults to client address)')
    
    args = parser.parse_args()
    
    if not args.command:
        parser.print_help()
        return 1
    
    # Execute command
    if args.command == 'parse':
        return parse_command(args)
    elif args.command == 'validate':
        return validate_command(args)
    elif args.command == 'demo':
        return demo_command(args)
    elif args.command == 'convert':
        return convert_command(args)
    elif args.command == 'send':
        return send_command(args)
    elif args.command == 'api':
        return api_command(args)
    elif args.command == 'list-conversions':
        return list_conversions_command(args)
    elif args.command == 'send-saved':
        return send_saved_command(args)
    
    return 0


if __name__ == '__main__':
    sys.exit(main())
