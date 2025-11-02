"""
Balance File Parser for Lynx Crypto Converter
Extracts numeric balance values from .docx/.dox files
"""

from docx import Document
import re
from decimal import Decimal
from typing import List, Dict, Optional
import os


class BalanceParser:
    """Parse balance information from document files"""
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.balances = []
        self._validate_file()
    
    def _validate_file(self) -> None:
        """Validate file exists and has correct extension"""
        if not os.path.exists(self.file_path):
            raise FileNotFoundError(f"File not found: {self.file_path}")
        
        ext = os.path.splitext(self.file_path)[1].lower()
        if ext not in ['.docx', '.dox']:
            raise ValueError(f"Invalid file type: {ext}. Expected .docx or .dox")
    
    def parse(self) -> List[Dict]:
        """
        Extract numeric balances from document
        
        Returns:
            List of dictionaries containing:
            - value: Decimal amount
            - currency_symbol: Optional currency symbol found
            - context: Surrounding text
            - line_number: Line position in document
        """
        try:
            doc = Document(self.file_path)
            
            for idx, para in enumerate(doc.paragraphs, 1):
                text = para.text.strip()
                if text:
                    numbers = self._extract_numbers(text, idx)
                    if numbers:
                        self.balances.extend(numbers)
            
            # Also check tables
            for table in doc.tables:
                for row_idx, row in enumerate(table.rows):
                    for cell in row.cells:
                        text = cell.text.strip()
                        if text:
                            numbers = self._extract_numbers(text, f"table-{row_idx}")
                            if numbers:
                                self.balances.extend(numbers)
            
            return self.balances
        
        except Exception as e:
            raise Exception(f"Error parsing document: {str(e)}")
    
    def _extract_numbers(self, text: str, line_ref) -> List[Dict]:
        """
        Extract numeric values with optional currency symbols
        
        Patterns matched:
        - $1,234.56
        - €1.234,56
        - 1234.56
        - 1,234.56
        """
        results = []
        
        # Pattern for currency symbols and numbers
        # Matches: $1,234.56 or €1.234,56 or 1234.56
        pattern = r'([$€£¥₹]?\s*\d{1,3}(?:[,.\s]\d{3})*(?:[.,]\d{2})?)'
        
        matches = re.finditer(pattern, text)
        
        for match in matches:
            raw_value = match.group(1)
            
            # Extract currency symbol if present
            currency_match = re.match(r'^([$€£¥₹])', raw_value)
            currency_symbol = currency_match.group(1) if currency_match else None
            
            # Clean number string
            num_str = re.sub(r'[$€£¥₹\s]', '', raw_value)
            
            # Handle different decimal separators
            # If comma is decimal separator (European format)
            if ',' in num_str and '.' in num_str:
                # Assume dot is thousands separator, comma is decimal
                if num_str.rindex(',') > num_str.rindex('.'):
                    num_str = num_str.replace('.', '').replace(',', '.')
                else:
                    num_str = num_str.replace(',', '')
            elif ',' in num_str:
                # Check if it's likely a decimal separator
                parts = num_str.split(',')
                if len(parts) == 2 and len(parts[1]) == 2:
                    num_str = num_str.replace(',', '.')
                else:
                    num_str = num_str.replace(',', '')
            
            try:
                value = Decimal(num_str)
                
                # Filter out unrealistic values (too small or suspiciously large)
                if value >= Decimal('0.01') and value <= Decimal('999999999999'):
                    results.append({
                        'value': float(value),  # Convert to float for JSON serialization
                        'value_decimal': str(value),  # Keep string representation
                        'currency_symbol': currency_symbol,
                        'original_text': match.group(0),
                        'context': text,
                        'line_ref': line_ref
                    })
            except (ValueError, Exception):
                # Skip invalid numbers
                continue
        
        return results
    
    def get_total(self) -> Decimal:
        """Calculate total of all extracted balances"""
        return sum(Decimal(b['value_decimal']) for b in self.balances)
    
    def get_summary(self) -> Dict:
        """Get summary statistics of parsed balances"""
        if not self.balances:
            return {
                'total_values_found': 0,
                'total_sum': 0,
                'min_value': 0,
                'max_value': 0,
                'avg_value': 0
            }
        
        values = [Decimal(b['value_decimal']) for b in self.balances]
        
        return {
            'total_values_found': len(self.balances),
            'total_sum': float(sum(values)),
            'min_value': float(min(values)),
            'max_value': float(max(values)),
            'avg_value': float(sum(values) / len(values))
        }
