#!/usr/bin/env python3
"""End-to-end conversion test"""

import sys
import os
import json
sys.path.insert(0, 'src')

from dotenv import load_dotenv
load_dotenv()

from src.converter import crypto_converter

print('=' * 60)
print('END-TO-END CONVERSION TEST')
print('=' * 60)

# First, create demo file if it doesn't exist
demo_file = 'src/demo_balances.docx'
if not os.path.exists(demo_file):
    print('\nCreating demo balance file...')
    from docx import Document

    doc = Document()
    doc.add_heading('Account Balances - November 2024', 0)
    doc.add_paragraph('Checking Account: $5,250.00')
    doc.add_paragraph('Savings Account: $12,800.50')
    doc.add_paragraph('Investment Portfolio: $45,000.00')
    doc.add_paragraph('Emergency Fund: $8,500.00')
    doc.add_paragraph('Crypto Wallet: $3,275.25')
    doc.save(demo_file)
    print('✓ Demo file created')

print('\n' + '=' * 60)
print('TEST 1: Full Balance Conversion')
print('=' * 60)

result = crypto_converter.convert_balances(demo_file)

if result.get('success'):
    print('\n✓ Conversion successful!')

    print(f'\nParsed {len(result["parsed_balances"])} balance entries')
    print(f'Total USD Amount: ${result["total_usd_amount"]:,.2f}')

    print('\nCurrent Exchange Rates (1 crypto = X USD):')
    for currency, rate in result['rates'].items():
        print(f'  1 {currency} = ${rate:,.2f} USD')

    print(f'\nConversions (How much crypto you can buy with ${result["total_usd_amount"]:,.2f}):')
    for currency, amount in result['conversions'].items():
        print(f'  {currency}: {amount:.8f}')

    print('\nWallet Information:')
    for currency, info in result['wallet_info'].items():
        status = '✓' if info['valid'] else '✗'
        addr = info['address'] if info['address'] else 'Not configured'
        print(f'  {currency} {status}:')
        print(f'    Address: {addr}')
        print(f'    Amount: {info["amount"]:.8f} {currency}')
else:
    print(f'\n✗ Conversion failed: {result.get("error")}')
    sys.exit(1)

print('\n' + '=' * 60)
print('TEST 2: Single Amount Conversion')
print('=' * 60)

test_amount = 1000.0
print(f'\nConverting ${test_amount:,.2f} USD to BTC...')

result2 = crypto_converter.convert_single_amount(test_amount, 'USD', 'BTC')

if result2.get('success'):
    print('\n✓ Single conversion successful!')
    print(f'\n{result2["calculation"]}')
    print(f'\nResult: {result2["converted_amount"]:.8f} BTC')
else:
    print(f'\n✗ Single conversion failed: {result2.get("error")}')
    sys.exit(1)

print('\n' + '=' * 60)
print('TEST 3: Portfolio Summary')
print('=' * 60)

result3 = crypto_converter.get_portfolio_summary(demo_file)

if result3.get('success'):
    print('\n✓ Portfolio summary generated!')

    ws = result3['wallet_summary']
    print('\nWallet Summary:')
    print(f'  Total Wallets: {ws["total_wallets"]}')
    print(f'  Valid Wallets: {ws["valid_wallets"]}')
    print(f'  Invalid Wallets: {ws["invalid_wallets"]}')
    print(f'  Missing Wallets: {ws["missing_wallets"]}')
else:
    print(f'\n✗ Portfolio summary failed: {result3.get("error")}')
    sys.exit(1)

print('\n' + '=' * 60)
print('✓ ALL TESTS PASSED!')
print('=' * 60)
print('\nMilestone 2 Core Functionality: WORKING')
print('Ready for API testing and client demo.')
